"""
Shared pytest fixtures for ExamOps tests.
"""

import io
import pytest
from docx import Document


def _make_doc(*paragraphs: str) -> Document:
    """Create a minimal python-docx Document with the given paragraph texts."""
    doc = Document()
    # Clear any default content (python-docx may or may not add a blank paragraph)
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


@pytest.fixture
def sample_doc():
    """A Document with typical pre-formatted exam paragraph patterns."""
    return _make_doc(
        "Q1) What is the capital of France?",
        "1a) Describe the geography.  [3 marks]",
        "1b) Explain the history.  (2 Marks)",
        "DATE :  January 2025",
        "COURSE CODE :  CS101",
    )


@pytest.fixture
def empty_doc():
    """A Document with no paragraphs."""
    doc = Document()
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)
    return doc


@pytest.fixture
def mock_template_rules():
    """Minimal template rules dict matching Azure AI Search output schema."""
    return {
        "header_text": "SOUTHERN UNIVERSITY COLLEGE",
        "footer_text": "Page",
        "margin_cm": {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.5},
        "numbering_style": "Q1. / (a) / (i)",
        "marks_format": "(N marks)",
    }


@pytest.fixture
def mock_validation_result():
    """A typical successful LLM validation result dict."""
    return {
        "compliance_score": 92.5,
        "category_scores": {
            "numbering": 95,
            "spacing": 90,
            "marks": 95,
            "header": 100,
            "indent": 85,
        },
        "issues_found": [
            "numbering style Q1) should be Q1.",
            "mark notation [3 marks] should be (3 marks)",
            "spacing: colon not padded in DATE line",
        ],
        "edge_cases": [],
        "math_expressions_preserved": True,
        "summary": "Document largely compliant with minor numbering and spacing fixes.",
        "fallback_mode": False,
    }


@pytest.fixture
def fallback_validation_result():
    """A validation result dict representing LLM timeout fallback."""
    return {
        "compliance_score": None,
        "category_scores": {},
        "issues_found": [],
        "edge_cases": [],
        "math_expressions_preserved": True,
        "summary": "LLM validation unavailable.",
        "fallback_mode": True,
        "error": "Connection timeout",
    }
