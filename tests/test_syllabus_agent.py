"""
Unit tests for src/agents/syllabus_agent/syllabus_agent.py

Covers:
- .docx text extraction calls LLM → returns clo_list/plo_list
- .pdf extraction path
- Malformed LLM JSON → returns empty lists (no crash)
- Empty file → returns empty lists
"""

import io
import json
import sys
import types
import pytest
from unittest.mock import AsyncMock, MagicMock, patch
from docx import Document


# ---------------------------------------------------------------------------
# Ensure openai stub is present (mirrors test_llm_client.py setup)
# ---------------------------------------------------------------------------

def _ensure_openai_stub():
    if "openai" not in sys.modules or not hasattr(sys.modules["openai"], "AsyncAzureOpenAI"):
        stub = types.ModuleType("openai")
        stub.AsyncAzureOpenAI = MagicMock
        stub.AsyncOpenAI = MagicMock

        class _RateLimitError(Exception):
            def __init__(self, message="", response=None, body=None):
                super().__init__(message)

        stub.RateLimitError = _RateLimitError
        sys.modules["openai"] = stub


_ensure_openai_stub()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_bytes(*paragraphs: str) -> bytes:
    """Create minimal .docx bytes with the given paragraph texts."""
    doc = Document()
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


SAMPLE_CLO_RESPONSE = json.dumps({
    "clo_list": ["CLO1: Understand AI concepts", "CLO2: Apply machine learning"],
    "plo_list": ["PLO1: Problem solving", "PLO2: Communication"],
})


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(autouse=True)
def _env_vars(monkeypatch):
    monkeypatch.setenv("AZURE_OPENAI_ENDPOINT", "https://fake.openai.azure.com/")
    monkeypatch.setenv("AZURE_OPENAI_KEY", "fake-key")
    monkeypatch.setenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4o-mini")
    monkeypatch.setenv("GITHUB_TOKEN", "fake-gh-token")


@pytest.fixture(autouse=True)
def _patch_session_store():
    """Suppress real Azure Table Storage calls."""
    with patch(
        "src.agents.syllabus_agent.syllabus_agent.SyllabusAgent._update_session",
        new=AsyncMock(),
    ):
        yield


# ---------------------------------------------------------------------------
# Test: .docx extraction calls LLM → returns clo_list / plo_list
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_process_docx_returns_clos_and_plos():
    """process() extracts CLO/PLO from a .docx file via LLM."""
    from src.agents.syllabus_agent.syllabus_agent import SyllabusAgent

    agent = SyllabusAgent()
    agent._llm.chat = AsyncMock(return_value=SAMPLE_CLO_RESPONSE)

    docx_bytes = _make_docx_bytes(
        "CLO1: Understand AI concepts",
        "CLO2: Apply machine learning",
        "PLO1: Problem solving",
    )

    result = await agent.process(docx_bytes, "syllabus.docx", "session-001")

    agent._llm.chat.assert_awaited_once()
    assert result["clo_list"] == [
        "CLO1: Understand AI concepts",
        "CLO2: Apply machine learning",
    ]
    assert result["plo_list"] == ["PLO1: Problem solving", "PLO2: Communication"]


# ---------------------------------------------------------------------------
# Test: .pdf extraction path
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_process_pdf_extraction_path():
    """process() calls _extract_text_from_pdf for .pdf filenames."""
    from src.agents.syllabus_agent import syllabus_agent as sa_module
    from src.agents.syllabus_agent.syllabus_agent import SyllabusAgent

    agent = SyllabusAgent()
    agent._llm.chat = AsyncMock(return_value=SAMPLE_CLO_RESPONSE)

    fake_pdf = b"%PDF-1.4 fake content with CLO information"

    with patch.object(sa_module, "_extract_text_from_pdf", return_value="CLO1: Test CLO"):
        result = await agent.process(fake_pdf, "syllabus.pdf", "session-002")

    assert "clo_list" in result
    assert "plo_list" in result
    agent._llm.chat.assert_awaited_once()


# ---------------------------------------------------------------------------
# Test: Malformed LLM JSON → returns empty lists (no crash)
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_malformed_llm_response_returns_empty_lists():
    """process() returns empty lists when LLM returns non-JSON content."""
    from src.agents.syllabus_agent.syllabus_agent import SyllabusAgent

    agent = SyllabusAgent()
    agent._llm.chat = AsyncMock(return_value="Not valid JSON at all!!!")

    docx_bytes = _make_docx_bytes("Some syllabus text with CLOs listed here.")

    result = await agent.process(docx_bytes, "syllabus.docx", "session-003")

    assert result == {"clo_list": [], "plo_list": []}


# ---------------------------------------------------------------------------
# Test: Empty file → returns empty lists
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_empty_file_returns_empty_lists():
    """process() returns empty lists without calling LLM when text is empty."""
    from src.agents.syllabus_agent.syllabus_agent import SyllabusAgent

    agent = SyllabusAgent()
    agent._llm.chat = AsyncMock(return_value=SAMPLE_CLO_RESPONSE)

    # A .docx with no paragraphs → empty text extraction
    empty_docx = _make_docx_bytes()

    result = await agent.process(empty_docx, "empty.docx", "session-004")

    assert result == {"clo_list": [], "plo_list": []}
    agent._llm.chat.assert_not_awaited()
