"""
Unit tests for DiffGeneratorAgent.

Tests cover:
  - generate_summary_stats: correct category counts from issues_found
  - _add_css_styling: injects CSS when missing, no-ops when already present
  - create_html_diff: returns dict with 'html_report' and 'summary_stats' keys
  - _extract_text_with_formatting: returns newline-joined paragraph text
"""

import pytest
from docx import Document

from src.agents.diff_generator.diff_generator import DiffGeneratorAgent, _DIFF_CSS


def _make_doc(*paragraphs: str) -> Document:
    doc = Document()
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


# ── generate_summary_stats ────────────────────────────────────────────────────

class TestGenerateSummaryStats:
    def setup_method(self):
        self.agent = DiffGeneratorAgent()

    def test_counts_numbering_fixes(self):
        validation = {
            "issues_found": ["numbering Q1) should be Q1.", "numbering style mismatch"],
            "compliance_score": 85.0,
        }
        doc = _make_doc("Q1. content")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["numbering_fixes"] == 2

    def test_counts_spacing_fixes(self):
        validation = {
            "issues_found": ["spacing: colon not padded", "colon missing space"],
            "compliance_score": 90.0,
        }
        doc = _make_doc("DATE : Jan")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["spacing_fixes"] == 2

    def test_counts_mark_fixes(self):
        validation = {
            "issues_found": ["mark notation should use parentheses"],
            "compliance_score": 88.0,
        }
        doc = _make_doc("content (3 marks)")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["mark_formatting_fixes"] == 1

    def test_counts_indentation_fixes(self):
        validation = {
            "issues_found": ["indent level incorrect for (a)", "indent missing"],
            "compliance_score": 80.0,
        }
        doc = _make_doc("(a) content")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["indentation_fixes"] == 2

    def test_total_changes_summed(self):
        validation = {
            "issues_found": [
                "numbering fix",
                "spacing fix",
                "mark fix",
                "indent fix",
            ],
            "compliance_score": 75.0,
        }
        doc = _make_doc("Q1. text")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["total_changes"] == stats["numbering_fixes"] + stats["spacing_fixes"] + \
               stats["mark_formatting_fixes"] + stats["indentation_fixes"] + \
               (1 if stats["header_footer_added"] else 0)

    def test_compliance_score_propagated(self):
        validation = {"issues_found": [], "compliance_score": 99.0}
        doc = _make_doc("Q1. text")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["compliance_score"] == 99.0

    def test_compliance_score_none_when_fallback(self):
        validation = {"issues_found": [], "compliance_score": None}
        doc = _make_doc("text")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["compliance_score"] is None

    def test_empty_issues(self):
        validation = {"issues_found": [], "compliance_score": 100.0}
        doc = _make_doc("Q1. Perfect document.")
        stats = self.agent.generate_summary_stats(doc, doc, validation)
        assert stats["numbering_fixes"] == 0
        assert stats["spacing_fixes"] == 0
        assert stats["mark_formatting_fixes"] == 0
        assert stats["indentation_fixes"] == 0


# ── _add_css_styling ──────────────────────────────────────────────────────────

class TestAddCssStyling:
    def setup_method(self):
        self.agent = DiffGeneratorAgent()

    def test_injects_css_when_missing(self):
        html = "<html><head></head><body></body></html>"
        result = self.agent._add_css_styling(html)
        assert "<style>" in result

    def test_no_op_when_style_already_present(self):
        html = "<html><head><style>body{}</style></head><body></body></html>"
        result = self.agent._add_css_styling(html)
        # Should not duplicate
        assert result.count("<style>") == 1

    def test_css_contains_diff_add_class(self):
        html = "<html><head></head><body></body></html>"
        result = self.agent._add_css_styling(html)
        assert "diff_add" in result

    def test_css_contains_diff_sub_class(self):
        html = "<html><head></head><body></body></html>"
        result = self.agent._add_css_styling(html)
        assert "diff_sub" in result


# ── create_html_diff ──────────────────────────────────────────────────────────

class TestCreateHtmlDiff:
    def setup_method(self):
        self.agent = DiffGeneratorAgent()

    def test_returns_html_and_stats_keys(self):
        original = _make_doc("Q1) What is x?", "[2 marks]")
        formatted = _make_doc("Q1.  What is x?", "(2 marks)")
        validation = {"issues_found": ["numbering", "mark fix"], "compliance_score": 90.0}

        result = self.agent.create_html_diff(original, formatted, validation)

        assert "html_report" in result
        assert "summary_stats" in result

    def test_html_report_is_string(self):
        doc = _make_doc("Q1. text")
        validation = {"issues_found": [], "compliance_score": 100.0}

        result = self.agent.create_html_diff(doc, doc, validation)

        assert isinstance(result["html_report"], str)
        assert len(result["html_report"]) > 0

    def test_html_report_contains_doctype(self):
        doc = _make_doc("Q1. text")
        validation = {"issues_found": [], "compliance_score": 100.0}

        result = self.agent.create_html_diff(doc, doc, validation)

        assert "<!DOCTYPE html>" in result["html_report"]

    def test_html_report_has_css(self):
        doc = _make_doc("Q1. text")
        validation = {"issues_found": [], "compliance_score": 100.0}

        result = self.agent.create_html_diff(doc, doc, validation)

        assert "<style>" in result["html_report"]

    def test_summary_stats_is_dict(self):
        doc = _make_doc("Q1. text")
        validation = {"issues_found": [], "compliance_score": 100.0}

        result = self.agent.create_html_diff(doc, doc, validation)

        assert isinstance(result["summary_stats"], dict)


# ── _extract_text_with_formatting ─────────────────────────────────────────────

class TestExtractText:
    def setup_method(self):
        self.agent = DiffGeneratorAgent()

    def test_joins_paragraphs_with_newlines(self):
        doc = _make_doc("First line", "Second line", "Third line")
        text = self.agent._extract_text_with_formatting(doc)
        assert "First line\nSecond line\nThird line" == text

    def test_empty_doc_returns_empty_string(self):
        doc = Document()
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        text = self.agent._extract_text_with_formatting(doc)
        assert text == ""
