"""
Unit tests for CoordinatorAgent.

Tests cover:
  - process_job success path (all agents mocked)
  - ERR_CORRUPTED_FILE — download_from_blob raises
  - ERR_TEMPLATE_NOT_FOUND — get_template_from_vectordb raises / returns empty
  - LLM timeout fallback — validation fallback_mode=True → status="partial"
  - ERR_STORAGE — save_outputs raises
"""

import pytest
from unittest.mock import AsyncMock, MagicMock, patch
from docx import Document

from src.agents.coordinator_agent.coordinator_agent import CoordinatorAgent, JobState


# ── Fixtures ─────────────────────────────────────────────────────────────────

def _make_mock_doc() -> Document:
    doc = Document()
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    doc.add_paragraph("Q1. What is x?")
    return doc


MOCK_TEMPLATE_RULES = {
    "header_text": "SOUTHERN UNIVERSITY COLLEGE",
    "margin_cm": {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.5},
}

MOCK_VALIDATION = {
    "compliance_score": 88.0,
    "issues_found": ["numbering fix"],
    "fallback_mode": False,
    "summary": "Good compliance.",
}

MOCK_FALLBACK_VALIDATION = {
    "compliance_score": None,
    "issues_found": [],
    "fallback_mode": True,
    "summary": "LLM unavailable.",
}

MOCK_OUTPUT_URLS = {
    "docx": "https://blob.example.com/output/job123.docx?sas=token",
    "html": "https://blob.example.com/output/job123-diff.html?sas=token",
}

MOCK_DIFF_RESULT = {
    "html_report": "<html>diff</html>",
    "summary_stats": {
        "numbering_fixes": 1,
        "spacing_fixes": 0,
        "mark_formatting_fixes": 0,
        "indentation_fixes": 0,
        "header_footer_added": True,
        "total_changes": 2,
        "compliance_score": 88.0,
    },
}


def _make_coordinator_with_mocks(
    *,
    download_return=None,
    download_raises=None,
    template_return=None,
    template_raises=None,
    format_return=None,
    diff_return=None,
    save_return=None,
    save_raises=None,
    onedrive_return="https://onedrive.example.com/link",
):
    """Build a CoordinatorAgent with all sub-agents replaced by mocks."""
    coord = CoordinatorAgent.__new__(CoordinatorAgent)

    # FileHandlerAgent mock
    fh = MagicMock()
    if download_raises:
        fh.download_from_blob = AsyncMock(side_effect=download_raises)
    else:
        fh.download_from_blob = AsyncMock(return_value=download_return or _make_mock_doc())

    if template_raises:
        fh.get_template_from_vectordb = AsyncMock(side_effect=template_raises)
    else:
        fh.get_template_from_vectordb = AsyncMock(
            return_value=template_return if template_return is not None else MOCK_TEMPLATE_RULES
        )

    if save_raises:
        fh.save_outputs = AsyncMock(side_effect=save_raises)
    else:
        fh.save_outputs = AsyncMock(return_value=save_return or MOCK_OUTPUT_URLS)

    fh.create_onedrive_link = AsyncMock(return_value=onedrive_return)

    # FormattingEngineAgent mock
    fe = MagicMock()
    formatted_doc = _make_mock_doc()
    fe.process_and_validate = AsyncMock(
        return_value=(formatted_doc, format_return or MOCK_VALIDATION)
    )

    # DiffGeneratorAgent mock
    dg = MagicMock()
    dg.create_html_diff = MagicMock(return_value=diff_return or MOCK_DIFF_RESULT)

    coord.file_handler = fh
    coord.formatting_engine = fe
    coord.diff_generator = dg

    return coord


# ── Success path ──────────────────────────────────────────────────────────────

class TestProcessJobSuccess:
    @pytest.mark.asyncio
    async def test_returns_status_success(self):
        coord = _make_coordinator_with_mocks()
        result = await coord.process_job("job-001", "user-abc", "https://blob/input.docx")
        assert result["status"] == "success"

    @pytest.mark.asyncio
    async def test_returns_compliance_score(self):
        coord = _make_coordinator_with_mocks()
        result = await coord.process_job("job-001", "user-abc", "https://blob/input.docx")
        assert result["compliance_score"] == 88.0

    @pytest.mark.asyncio
    async def test_returns_formatted_url(self):
        coord = _make_coordinator_with_mocks()
        result = await coord.process_job("job-001", "user-abc", "https://blob/input.docx")
        assert result["formatted_url"] == MOCK_OUTPUT_URLS["docx"]

    @pytest.mark.asyncio
    async def test_returns_diff_url(self):
        coord = _make_coordinator_with_mocks()
        result = await coord.process_job("job-001", "user-abc", "https://blob/input.docx")
        assert result["diff_url"] == MOCK_OUTPUT_URLS["html"]

    @pytest.mark.asyncio
    async def test_returns_onedrive_link(self):
        coord = _make_coordinator_with_mocks()
        result = await coord.process_job("job-001", "user-abc", "https://blob/input.docx")
        assert "onedrive.example.com" in result["onedrive_link"]

    @pytest.mark.asyncio
    async def test_no_error_on_success(self):
        coord = _make_coordinator_with_mocks()
        result = await coord.process_job("job-001", "user-abc", "https://blob/input.docx")
        assert result["error"] is None

    @pytest.mark.asyncio
    async def test_summary_contains_changes(self):
        coord = _make_coordinator_with_mocks()
        result = await coord.process_job("job-001", "user-abc", "https://blob/input.docx")
        assert "change" in result["summary"].lower()


# ── ERR_CORRUPTED_FILE ────────────────────────────────────────────────────────

class TestCorruptedFile:
    @pytest.mark.asyncio
    async def test_status_is_failed(self):
        coord = _make_coordinator_with_mocks(download_raises=ValueError("bad docx"))
        result = await coord.process_job("job-002", "user-abc", "https://blob/bad.docx")
        assert result["status"] == "failed"

    @pytest.mark.asyncio
    async def test_error_code_in_result(self):
        coord = _make_coordinator_with_mocks(download_raises=ValueError("bad docx"))
        result = await coord.process_job("job-002", "user-abc", "https://blob/bad.docx")
        assert "ERR_CORRUPTED_FILE" in result["error"]

    @pytest.mark.asyncio
    async def test_no_formatted_url(self):
        coord = _make_coordinator_with_mocks(download_raises=ValueError("bad docx"))
        result = await coord.process_job("job-002", "user-abc", "https://blob/bad.docx")
        assert result["formatted_url"] == ""


# ── ERR_TEMPLATE_NOT_FOUND ────────────────────────────────────────────────────

class TestTemplateNotFound:
    @pytest.mark.asyncio
    async def test_raises_returns_error_code(self):
        coord = _make_coordinator_with_mocks(
            template_raises=RuntimeError("search unavailable")
        )
        result = await coord.process_job("job-003", "user-abc", "https://blob/input.docx")
        assert result["status"] == "failed"
        assert "ERR_TEMPLATE_NOT_FOUND" in result["error"]

    @pytest.mark.asyncio
    async def test_empty_rules_returns_error_code(self):
        """get_template_from_vectordb returning {} should also trigger ERR_TEMPLATE_NOT_FOUND."""
        coord = _make_coordinator_with_mocks(template_return={})
        result = await coord.process_job("job-003b", "user-abc", "https://blob/input.docx")
        assert result["status"] == "failed"
        assert "ERR_TEMPLATE_NOT_FOUND" in result["error"]


# ── LLM timeout → partial ─────────────────────────────────────────────────────

class TestLLMTimeoutFallback:
    @pytest.mark.asyncio
    async def test_status_is_partial(self):
        coord = _make_coordinator_with_mocks(format_return=MOCK_FALLBACK_VALIDATION)
        result = await coord.process_job("job-004", "user-abc", "https://blob/input.docx")
        assert result["status"] == "partial"

    @pytest.mark.asyncio
    async def test_compliance_score_is_none(self):
        coord = _make_coordinator_with_mocks(format_return=MOCK_FALLBACK_VALIDATION)
        result = await coord.process_job("job-004", "user-abc", "https://blob/input.docx")
        assert result["compliance_score"] is None

    @pytest.mark.asyncio
    async def test_summary_mentions_llm_timeout(self):
        coord = _make_coordinator_with_mocks(format_return=MOCK_FALLBACK_VALIDATION)
        result = await coord.process_job("job-004", "user-abc", "https://blob/input.docx")
        assert "LLM" in result["summary"] or "rule-based" in result["summary"]

    @pytest.mark.asyncio
    async def test_formatted_url_still_returned(self):
        """Even in partial mode, a formatted document URL should still be present."""
        coord = _make_coordinator_with_mocks(format_return=MOCK_FALLBACK_VALIDATION)
        result = await coord.process_job("job-004", "user-abc", "https://blob/input.docx")
        assert result["formatted_url"] != ""


# ── ERR_STORAGE ───────────────────────────────────────────────────────────────

class TestStorageError:
    @pytest.mark.asyncio
    async def test_status_is_failed(self):
        coord = _make_coordinator_with_mocks(
            save_raises=OSError("blob storage connection refused")
        )
        result = await coord.process_job("job-005", "user-abc", "https://blob/input.docx")
        assert result["status"] == "failed"

    @pytest.mark.asyncio
    async def test_error_code_in_result(self):
        coord = _make_coordinator_with_mocks(
            save_raises=OSError("blob storage connection refused")
        )
        result = await coord.process_job("job-005", "user-abc", "https://blob/input.docx")
        assert "ERR_STORAGE" in result["error"]


# ── JobState ──────────────────────────────────────────────────────────────────

class TestJobState:
    def test_initial_status_pending(self):
        job = JobState(job_id="j1", user_id="u1", file_url="https://blob/x.docx")
        assert job.status == "pending"

    def test_update_status_changes_status(self):
        job = JobState(job_id="j1", user_id="u1", file_url="https://blob/x.docx")
        job.update_status("downloading")
        assert job.status == "downloading"

    def test_update_status_sets_error(self):
        job = JobState(job_id="j1", user_id="u1", file_url="https://blob/x.docx")
        job.update_status("failed", error="ERR_STORAGE")
        assert job.error == "ERR_STORAGE"

    def test_update_status_updates_timestamp(self):
        job = JobState(job_id="j1", user_id="u1", file_url="https://blob/x.docx")
        original_ts = job.updated_at
        job.update_status("formatting")
        assert job.updated_at >= original_ts
