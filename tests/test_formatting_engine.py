"""
Unit tests for FormattingEngineAgent — RuleBasedFormatter and LLMValidator.

Tests cover:
  - _fix_numbering: Q1) → Q1.  |  1a) → (a)
  - _format_marks: [3 marks] → (3 marks)  |  (2 Marks) → (2 marks)
  - _enforce_spacing: "DATE :" → "DATE : "
  - m:oMath guard: math paragraphs are skipped
  - LLMValidator fallback on timeout/exception
  - FormattingEngineAgent.process_and_validate integration
"""

import pytest
import pytest_asyncio
from unittest.mock import AsyncMock, MagicMock, patch
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from src.agents.formatting_engine.formatting_engine import (
    RuleBasedFormatter,
    LLMValidator,
    FormattingEngineAgent,
)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _make_para_doc(text: str) -> tuple:
    """Return (doc, para) with a single paragraph containing text."""
    doc = Document()
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    para = doc.add_paragraph(text)
    return doc, para


def _inject_omath(para) -> None:
    """Inject a fake m:oMath element into the first run of para."""
    run = para.runs[0] if para.runs else para.add_run("x")
    omath_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    omath_elem = etree.SubElement(run._element, f"{{{omath_ns}}}oMath")


# ── RuleBasedFormatter._fix_numbering ────────────────────────────────────────

class TestFixNumbering:
    def setup_method(self):
        self.fmt = RuleBasedFormatter()

    def test_q_bracket_to_dot(self):
        """Q1) should become Q1.  """
        _, para = _make_para_doc("Q1) What is 2+2?")
        self.fmt._fix_numbering(para)
        assert para.text.startswith("Q1.")

    def test_q_number_preserves_rest_of_text(self):
        """Text after the numbering marker should be preserved."""
        _, para = _make_para_doc("Q2) Calculate the value of x.")
        self.fmt._fix_numbering(para)
        assert "Calculate the value of x." in para.text

    def test_alpha_sub_bracket(self):
        """1a) should become (a) """
        _, para = _make_para_doc("1a) Explain the concept.")
        self.fmt._fix_numbering(para)
        assert para.text.startswith("(a)")

    def test_alpha_sub_preserves_letter(self):
        """Letter extracted should match input sub-question letter."""
        _, para = _make_para_doc("1c) Compare two approaches.")
        self.fmt._fix_numbering(para)
        assert para.text.startswith("(c)")

    def test_no_change_for_already_correct(self):
        """Paragraphs already in correct format should not be modified."""
        _, para = _make_para_doc("Q1.  What is the capital of France?")
        original = para.text
        self.fmt._fix_numbering(para)
        assert para.text == original

    def test_empty_paragraph_skipped(self):
        _, para = _make_para_doc("")
        self.fmt._fix_numbering(para)
        assert para.text == ""

    def test_math_paragraph_skipped(self):
        """A paragraph containing m:oMath must not be modified."""
        _, para = _make_para_doc("Q1) ∫f(x)dx = F(x) + C")
        _inject_omath(para)
        original = para.text
        # Process via the public interface so the guard fires
        fmt = RuleBasedFormatter()
        doc = Document()
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        # Add para directly — simulate single-paragraph doc
        doc._body._body.append(para._element)
        fmt.process(doc, {})
        assert para.text == original


# ── RuleBasedFormatter._format_marks ─────────────────────────────────────────

class TestFormatMarks:
    def setup_method(self):
        self.fmt = RuleBasedFormatter()

    def test_square_bracket_marks(self):
        """[3 marks] → (3 marks)"""
        _, para = _make_para_doc("Describe the process.  [3 marks]")
        self.fmt._format_marks(para)
        assert "(3 marks)" in para.text
        assert "[3 marks]" not in para.text

    def test_capital_marks(self):
        """(2 Marks) → (2 marks)"""
        _, para = _make_para_doc("Explain briefly.  (2 Marks)")
        self.fmt._format_marks(para)
        assert "(2 marks)" in para.text

    def test_singular_mark(self):
        """[1 mark] → (1 mark) — singular form."""
        _, para = _make_para_doc("State one example.  [1 mark]")
        self.fmt._format_marks(para)
        assert "(1 mark)" in para.text
        assert "marks" not in para.text

    def test_no_marks_unchanged(self):
        _, para = _make_para_doc("What is the capital of France?")
        original = para.text
        self.fmt._format_marks(para)
        assert para.text == original


# ── RuleBasedFormatter._enforce_spacing ──────────────────────────────────────

class TestEnforceSpacing:
    def setup_method(self):
        self.fmt = RuleBasedFormatter()

    def test_colon_space_added(self):
        """'DATE :' should become 'DATE : '"""
        _, para = _make_para_doc("DATE : January 2025")
        self.fmt._enforce_spacing(para)
        # After normalisation multiple spaces around colon become single spaces
        assert " : " in para.text

    def test_colon_no_trailing_space(self):
        """'DATE :January' (no trailing space) gets a space after."""
        _, para = _make_para_doc("COURSE CODE :CS101")
        self.fmt._enforce_spacing(para)
        assert " : " in para.text

    def test_no_colon_unchanged(self):
        _, para = _make_para_doc("No colon here at all")
        original = para.text
        self.fmt._enforce_spacing(para)
        assert para.text == original


# ── LLMValidator fallback ─────────────────────────────────────────────────────

class TestLLMValidatorFallback:
    @pytest.mark.asyncio
    async def test_fallback_on_exception(self, sample_doc, mock_template_rules):
        """LLMValidator must return fallback dict when _call_llm raises."""
        validator = LLMValidator()
        with patch.object(validator, "_call_llm", side_effect=TimeoutError("timeout")):
            result = await validator.validate(sample_doc, sample_doc, mock_template_rules)

        assert result["fallback_mode"] is True
        assert result["compliance_score"] is None
        assert "error" in result

    @pytest.mark.asyncio
    async def test_fallback_on_json_error(self, sample_doc, mock_template_rules):
        """LLMValidator must return fallback when LLM returns non-JSON."""
        validator = LLMValidator()
        with patch.object(validator, "_call_llm", side_effect=ValueError("bad json")):
            result = await validator.validate(sample_doc, sample_doc, mock_template_rules)

        assert result["fallback_mode"] is True


# ── FormattingEngineAgent integration ────────────────────────────────────────

class TestFormattingEngineAgent:
    @pytest.mark.asyncio
    async def test_process_and_validate_success(self, sample_doc, mock_template_rules, mock_validation_result):
        """process_and_validate should return (doc, validation_dict) on success."""
        agent = FormattingEngineAgent()
        with patch.object(agent.llm_validator, "validate", new=AsyncMock(return_value=mock_validation_result)):
            formatted_doc, validation = await agent.process_and_validate(sample_doc, mock_template_rules)

        assert isinstance(formatted_doc, Document)
        assert validation["compliance_score"] == 92.5
        assert validation["fallback_mode"] is False

    @pytest.mark.asyncio
    async def test_process_and_validate_llm_fallback(self, sample_doc, mock_template_rules, fallback_validation_result):
        """process_and_validate should still return a Document when LLM falls back."""
        agent = FormattingEngineAgent()
        with patch.object(agent.llm_validator, "validate", new=AsyncMock(return_value=fallback_validation_result)):
            formatted_doc, validation = await agent.process_and_validate(sample_doc, mock_template_rules)

        assert isinstance(formatted_doc, Document)
        assert validation["fallback_mode"] is True
        assert validation["compliance_score"] is None

    @pytest.mark.asyncio
    async def test_numbering_applied_end_to_end(self, mock_template_rules, fallback_validation_result):
        """Q1) paragraphs should be converted to Q1. after process_and_validate."""
        doc = Document()
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        doc.add_paragraph("Q1) Define machine learning.")

        agent = FormattingEngineAgent()
        with patch.object(agent.llm_validator, "validate", new=AsyncMock(return_value=fallback_validation_result)):
            formatted_doc, _ = await agent.process_and_validate(doc, mock_template_rules)

        first_para = formatted_doc.paragraphs[0].text
        assert first_para.startswith("Q1.")
