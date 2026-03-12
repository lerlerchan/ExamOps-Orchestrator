"""
ModerationFormAgent — fills the AARO-FM-030 moderation form .docx template
with session questions, CLO/PLO mapping, and marks.

Environment variables required:
    AZURE_STORAGE_CONNECTION_STRING
    BLOB_CONTAINER_OUTPUT       (default: examops-output)
    BLOB_CONTAINER_TEMPLATES    (default: examops-templates)
    MODERATION_TEMPLATE_BLOB    (default: AARO-FM-030.docx)
"""

import io
import logging
import os
from datetime import datetime, timezone

logger = logging.getLogger(__name__)

CONTAINER_OUTPUT = os.getenv("BLOB_CONTAINER_OUTPUT", "examops-output")
CONTAINER_TEMPLATES = os.getenv("BLOB_CONTAINER_TEMPLATES", "examops-templates")
MODERATION_TEMPLATE_BLOB = os.getenv("MODERATION_TEMPLATE_BLOB", "AARO-FM-030.docx")


class ModerationFormAgent:
    """
    Fills the AARO-FM-030 moderation form with questions from the session.

    The template is expected to contain a table with columns:
        No. | Question | CLO | PLO | Marks

    If the template is unavailable, a simple fallback .docx is generated.
    """

    def __init__(self) -> None:
        self._connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")

    async def fill_form(self, session_id: str) -> str:
        """
        Fill the moderation form for a session and upload to Blob Storage.

        Args:
            session_id: Session identifier.

        Returns:
            SAS URL of the uploaded moderation form .docx.
        """
        from src.session.session_store import SessionStore

        store = SessionStore()
        session = store.get_session(session_id)
        if not session:
            raise ValueError(f"Session not found: {session_id}")

        doc = await self._build_document(session)

        # Save to buffer and upload
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        blob_name = f"{session_id}_AARO-FM-030.docx"
        url = await self._upload(blob_name, buf.getvalue())

        # Update session
        session.moderation_form_url = url
        store.update_session(session)
        logger.info("Moderation form uploaded for session %s", session_id)
        return url

    async def _build_document(self, session):
        """Build the AARO-FM-030 document from session data."""
        from docx import Document

        # Try to load the template from Blob Storage
        doc = await self._load_template()

        if doc is None:
            # Generate a clean fallback document
            doc = Document()

        # Title block
        doc.add_heading("MODERATION FORM — AARO-FM-030", level=1)
        doc.add_paragraph(f"Session ID: {session.session_id}")
        doc.add_paragraph(f"Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
        doc.add_paragraph("")

        # CLO/PLO summary
        if session.clo_list:
            doc.add_heading("Course Learning Outcomes (CLOs)", level=2)
            for i, clo in enumerate(session.clo_list, 1):
                doc.add_paragraph(f"{i}. {clo}", style="List Number")

        if session.plo_list:
            doc.add_heading("Programme Learning Outcomes (PLOs)", level=2)
            for i, plo in enumerate(session.plo_list, 1):
                doc.add_paragraph(f"{i}. {plo}", style="List Number")

        doc.add_paragraph("")

        # Questions table
        doc.add_heading("Question Moderation Table", level=2)
        questions = session.questions or []

        if questions:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            # Header row
            hdr = table.rows[0].cells
            hdr[0].text = "No."
            hdr[1].text = "Question"
            hdr[2].text = "CLO"
            hdr[3].text = "PLO"
            hdr[4].text = "Marks"

            for hdr_cell in hdr:
                for para in hdr_cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            # Data rows
            for i, q in enumerate(questions, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                # Split multiline question text into separate paragraphs in the cell
                q_text = q.get("text", "")
                q_lines = [ln for ln in q_text.split("\n") if ln.strip()]
                if q_lines:
                    row[1].paragraphs[0].text = q_lines[0]
                    for extra_line in q_lines[1:]:
                        row[1].add_paragraph(extra_line)
                else:
                    row[1].text = q_text
                row[2].text = q.get("clo", "")
                row[3].text = q.get("plo", "")
                row[4].text = str(q.get("marks", ""))
        else:
            doc.add_paragraph("(No questions added to this session.)")

        return doc

    async def _load_template(self):
        """Attempt to load AARO-FM-030.docx template from Blob Storage."""
        try:
            from azure.storage.blob import BlobServiceClient
            from docx import Document

            service_client = BlobServiceClient.from_connection_string(
                self._connection_string
            )
            container_client = service_client.get_container_client(CONTAINER_TEMPLATES)
            blob_client = container_client.get_blob_client(MODERATION_TEMPLATE_BLOB)
            data = blob_client.download_blob().readall()
            return Document(io.BytesIO(data))
        except Exception as exc:
            logger.warning("Template not found in Blob Storage, using fallback: %s", exc)
            return None

    async def _upload(self, blob_name: str, data: bytes) -> str:
        """Upload bytes to examops-output and return a SAS URL."""
        from azure.storage.blob import (
            BlobServiceClient,
            generate_blob_sas,
            BlobSasPermissions,
            ContentSettings,
        )
        from datetime import timedelta

        service_client = BlobServiceClient.from_connection_string(
            self._connection_string
        )
        container_client = service_client.get_container_client(CONTAINER_OUTPUT)
        container_client.upload_blob(
            blob_name,
            data,
            overwrite=True,
            content_settings=ContentSettings(
                content_type=(
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                )
            ),
        )
        sas_token = generate_blob_sas(
            account_name=service_client.account_name,
            container_name=CONTAINER_OUTPUT,
            blob_name=blob_name,
            account_key=service_client.credential.account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.now(timezone.utc) + timedelta(hours=8),
        )
        return (
            f"https://{service_client.account_name}.blob.core.windows.net"
            f"/{CONTAINER_OUTPUT}/{blob_name}?{sas_token}"
        )
