"""
Azure Function HTTP trigger — POST /api/export-questions

Accepts JSON: {session_id}
Returns:      {download_url}  — .docx of all session questions for Step 5.
"""

import io
import json
import logging
from datetime import datetime, timedelta, timezone

import azure.functions as func

logger = logging.getLogger(__name__)

_CORS_HEADERS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type, X-Session-ID",
}

CONTAINER_OUTPUT = "examops-output"


async def main(req: func.HttpRequest) -> func.HttpResponse:
    if req.method == "OPTIONS":
        return func.HttpResponse(status_code=204, headers=_CORS_HEADERS)

    # ── Parse body ─────────────────────────────────────────────────────────
    try:
        body = req.get_json()
    except Exception:
        body = {}

    session_id = (
        body.get("session_id")
        or req.headers.get("X-Session-ID")
        or req.params.get("session_id")
    )

    if not session_id:
        return func.HttpResponse(
            json.dumps({"error": "Provide 'session_id'."}),
            status_code=400,
            mimetype="application/json",
            headers=_CORS_HEADERS,
        )

    # ── Load session ───────────────────────────────────────────────────────
    try:
        from src.session.session_store import SessionStore
        store = SessionStore()
        session = store.get_session(session_id)
        if not session:
            return func.HttpResponse(
                json.dumps({"error": f"Session not found: {session_id}"}),
                status_code=404,
                mimetype="application/json",
                headers=_CORS_HEADERS,
            )
    except Exception as exc:
        logger.exception("Session load failed")
        return func.HttpResponse(
            json.dumps({"error": f"Session error: {exc}"}),
            status_code=500,
            mimetype="application/json",
            headers=_CORS_HEADERS,
        )

    # ── Save questions from body if provided ───────────────────────────────
    questions_from_body = body.get("questions")
    if questions_from_body and isinstance(questions_from_body, list):
        session.questions = questions_from_body
        try:
            store.update_session(session)
        except Exception as exc:
            logger.warning("Could not persist questions to session: %s", exc)

    # ── Build .docx ────────────────────────────────────────────────────────
    try:
        doc = _build_questions_doc(session)
    except Exception as exc:
        logger.exception("Document build failed")
        return func.HttpResponse(
            json.dumps({"error": f"Document build failed: {exc}"}),
            status_code=500,
            mimetype="application/json",
            headers=_CORS_HEADERS,
        )

    # ── Upload to Blob ─────────────────────────────────────────────────────
    try:
        download_url = await _upload_doc(session_id, doc)
    except Exception as exc:
        logger.exception("Blob upload failed for export_questions")
        return func.HttpResponse(
            json.dumps({"error": f"Upload failed: {exc}"}),
            status_code=500,
            mimetype="application/json",
            headers=_CORS_HEADERS,
        )

    return func.HttpResponse(
        json.dumps({"session_id": session_id, "download_url": download_url}),
        status_code=200,
        mimetype="application/json",
        headers=_CORS_HEADERS,
    )


def _build_questions_doc(session):
    """Build a .docx listing all curated exam questions."""
    from docx import Document

    doc = Document()
    doc.add_heading("Exam Questions Export", level=1)
    doc.add_paragraph(f"Session: {session.session_id}")
    doc.add_paragraph(f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    questions = session.questions or []
    if not questions:
        doc.add_paragraph("No questions added to this session.")
        return doc

    for i, q in enumerate(questions, 1):
        doc.add_heading(f"Q{i}.", level=2)
        # Split on newlines so sub-parts each get their own paragraph
        raw_text = q.get("text", "(no text)")
        lines = [ln for ln in raw_text.split("\n") if ln.strip()]
        for line in (lines if lines else [raw_text]):
            doc.add_paragraph(line)
        meta_parts = []
        if q.get("clo"):
            meta_parts.append(f"CLO: {q['clo']}")
        if q.get("marks"):
            meta_parts.append(f"({q['marks']} marks)")
        if meta_parts:
            p = doc.add_paragraph(" | ".join(meta_parts))
            for run in p.runs:
                run.italic = True
        doc.add_paragraph("")

    return doc


async def _upload_doc(session_id: str, doc) -> str:
    """Upload .docx to blob and return SAS URL."""
    import os
    from azure.storage.blob import (
        BlobServiceClient,
        generate_blob_sas,
        BlobSasPermissions,
    )

    connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    blob_name = f"{session_id}_questions.docx"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    service_client = BlobServiceClient.from_connection_string(connection_string)
    container_client = service_client.get_container_client(CONTAINER_OUTPUT)
    from azure.storage.blob import ContentSettings as _CS
    container_client.upload_blob(
        blob_name,
        buf.getvalue(),
        overwrite=True,
        content_settings=_CS(
            content_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
        ),
    )

    sas_token = generate_blob_sas(
        account_name=service_client.account_name,
        container_name=CONTAINER_OUTPUT,
        blob_name=blob_name,
        account_key=service_client.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.now(timezone.utc) + timedelta(hours=8),
    )

    return (
        f"https://{service_client.account_name}.blob.core.windows.net"
        f"/{CONTAINER_OUTPUT}/{blob_name}?{sas_token}"
    )
