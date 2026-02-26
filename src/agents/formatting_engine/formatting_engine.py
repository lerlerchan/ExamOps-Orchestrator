"""
FormattingEngineAgent — two-layer hybrid document formatter.

Layer 1 — RuleBasedFormatter  (deterministic, python-docx)
Layer 2 — LLMValidator        (GPT-4o-mini via Azure AI Foundry)

The two layers are intentionally kept as separate classes and must NOT be merged.
"""

import json
import logging
import os
import re
from typing import Optional

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ── Constants ────────────────────────────────────────────────────────────────

# Indentation levels (Tab 0.5cm × multiplier)
INDENT_L1 = Cm(0.0)   # Q-level
INDENT_L2 = Cm(1.5)   # (a) level  — 0.5 × 3
INDENT_L3 = Cm(3.0)   # (i) level  — 0.5 × 6

# Numbering regex patterns
RE_Q_NUMBER = re.compile(r"^Q(\d+)[)]\s*")       # Q1)  → Q1.
RE_ALPHA_SUB = re.compile(r"^\d+[a-z][)]\s*")    # 1a)  → (a)
RE_ROMAN_SUB = re.compile(r"^\d+[.]\d+[.]\d+")   # 1.a.i → (i) style

# Marks pattern: ensure "(3 marks)" format
RE_MARKS = re.compile(r"\[(\d+)\s*marks?\]|\((\d+)\s*Marks?\)", re.IGNORECASE)

# Colon spacing: "DATE :" → "DATE : " (ensure space after colon)
RE_COLON = re.compile(r"\s+:\s*")


# ── Layer 1 — Rule-Based Formatter ──────────────────────────────────────────


class RuleBasedFormatter:
    """
    Deterministic python-docx transforms for exam paper compliance.

    Rules applied (in order):
        1. Header / footer injection
        2. Page margin standardisation
        3. Numbering correction  (Q1) → Q1.  |  1a) → (a)  |  1.a.i → (i))
        4. Marks notation        [3 marks] / (3 Marks) → (3 marks)
        5. Colon spacing         "DATE :" → "DATE : "
        6. Indentation           L1=0cm / L2=1.5cm / L3=3.0cm

    IMPORTANT: Any paragraph run whose XML contains ``m:oMath`` is skipped —
    math expression content is NEVER reformatted.
    """

    def process(self, doc: Document, template_rules: dict) -> Document:
        """
        Apply all rule-based transforms to ``doc`` in-place and return it.

        Args:
            doc:            python-docx Document to format.
            template_rules: Template rules dict from Azure AI Search.

        Returns:
            The same Document object, modified in-place.
        """
        self._apply_header_footer(doc, template_rules)
        self._standardize_margins(doc, template_rules)
        for para in doc.paragraphs:
            if self._contains_math(para):
                continue
            self._fix_numbering(para)
            self._format_marks(para)
            self._enforce_spacing(para)
            self._fix_indentation(para)
        return doc

    # ── Private transforms ───────────────────────────────────────────────────

    def _apply_header_footer(self, doc: Document, template_rules: dict) -> None:
        """Inject or overwrite the first-section header and footer text."""
        section = doc.sections[0]
        header_text = template_rules.get(
            "header_text", "SOUTHERN UNIVERSITY COLLEGE"
        )
        footer_text = template_rules.get("footer_text", "")

        # Header
        header = section.header
        if not header.paragraphs:
            header.add_paragraph(header_text)
        else:
            header.paragraphs[0].clear()
            header.paragraphs[0].add_run(header_text)

        # Footer (page number placeholder)
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph(footer_text or "Page")
        else:
            footer.paragraphs[0].clear()
            footer.paragraphs[0].add_run(footer_text or "Page")

    def _standardize_margins(self, doc: Document, template_rules: dict) -> None:
        """Set page margins from template_rules or institutional defaults."""
        margin_cfg = template_rules.get(
            "margin_cm",
            {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.5},
        )
        for section in doc.sections:
            section.top_margin = Cm(margin_cfg.get("top", 2.5))
            section.bottom_margin = Cm(margin_cfg.get("bottom", 2.5))
            section.left_margin = Cm(margin_cfg.get("left", 3.0))
            section.right_margin = Cm(margin_cfg.get("right", 2.5))

    def _fix_numbering(self, para) -> None:
        """
        Correct question/sub-question numbering styles.

        Conversions:
            Q1)   → Q1.
            1a)   → (a)
            1.a.i → (a)(i)  [best-effort regex normalisation]
        """
        text = para.text
        if not text.strip():
            return

        new_text: Optional[str] = None

        if RE_Q_NUMBER.match(text):
            new_text = RE_Q_NUMBER.sub(
                lambda m: f"Q{m.group(1)}.  ", text
            )
        elif RE_ALPHA_SUB.match(text):
            letter = re.match(r"^\d+([a-z])", text).group(1)
            new_text = re.sub(r"^\d+[a-z][)]\s*", f"({letter}) ", text)

        if new_text and new_text != text:
            self._replace_paragraph_text(para, new_text)

    def _format_marks(self, para) -> None:
        """
        Normalise marks notation to ``(N marks)`` / ``(1 mark)``.

        Handles:
            [3 marks]  → (3 marks)
            (3 Marks)  → (3 marks)
            [1 mark]   → (1 mark)
        """
        text = para.text

        def _normalise(match):
            n = int(match.group(1) or match.group(2))
            unit = "mark" if n == 1 else "marks"
            return f"({n} {unit})"

        new_text = RE_MARKS.sub(_normalise, text)
        if new_text != text:
            self._replace_paragraph_text(para, new_text)

    def _enforce_spacing(self, para) -> None:
        """
        Ensure a single space on both sides of a colon in header-style lines.

        Example: ``"DATE :"`` → ``"DATE : "``
        """
        text = para.text
        new_text = RE_COLON.sub(" : ", text)
        if new_text != text:
            self._replace_paragraph_text(para, new_text)

    def _fix_indentation(self, para) -> None:
        """
        Apply three-level indentation based on paragraph numbering level.

        Level detection:
            - Text starts with ``Q``  → L1 (no indent)
            - Text starts with ``(`` + single letter → L2 (1.5 cm)
            - Text starts with ``(`` + roman / digit  → L3 (3.0 cm)
        """
        text = para.text.lstrip()
        pf = para.paragraph_format

        if re.match(r"^Q\d+\.", text):
            pf.left_indent = INDENT_L1
        elif re.match(r"^\([a-z]\)", text):
            pf.left_indent = INDENT_L2
        elif re.match(r"^\([ivxlcdm]+\)|\(\d+\)", text, re.IGNORECASE):
            pf.left_indent = INDENT_L3

    # ── Utilities ────────────────────────────────────────────────────────────

    @staticmethod
    def _contains_math(para) -> bool:
        """Return True if any run in this paragraph contains an m:oMath element."""
        omath_tag = qn("m:oMath")
        for run in para.runs:
            if run._element.find(f".//{omath_tag}") is not None:
                return True
        return False

    @staticmethod
    def _replace_paragraph_text(para, new_text: str) -> None:
        """
        Replace the full text of a paragraph while preserving the first run's
        formatting.  Clears all runs and writes new_text into the first run.
        """
        if not para.runs:
            para.add_run(new_text)
            return
        first_run = para.runs[0]
        # Clear subsequent runs
        for run in para.runs[1:]:
            run.text = ""
        first_run.text = new_text


# ── Layer 2 — LLM Validator ─────────────────────────────────────────────────


class LLMValidator:
    """
    GPT-4o-mini validation layer (Azure AI Foundry).

    Runs after RuleBasedFormatter to handle:
    - Ambiguous numbering edge cases
    - Verification that m:oMath content was preserved
    - Overall compliance scoring (0–100%)

    On timeout or any exception, returns a fallback dict with
    ``compliance_score=None`` and ``fallback_mode=True``.
    """

    SYSTEM_PROMPT = (
        "You are an exam paper compliance checker for Southern University College. "
        "Evaluate the formatted document against the template rules and return a "
        "JSON object with the following keys: compliance_score (0-100), "
        "category_scores (dict), issues_found (list of str), "
        "edge_cases (list of str), math_expressions_preserved (bool), summary (str). "
        "Respond with JSON only — no markdown fences."
    )

    def __init__(self) -> None:
        self._endpoint = os.getenv("AZURE_FOUNDRY_ENDPOINT")
        self._api_key = os.getenv("AZURE_FOUNDRY_KEY")

    async def validate(
        self,
        original: Document,
        formatted: Document,
        template_rules: dict,
    ) -> dict:
        """
        Send both document texts to GPT-4o-mini and return a compliance dict.

        Args:
            original:       Original (unformatted) Document.
            formatted:      Formatted Document after Layer 1.
            template_rules: Template rules dict used during formatting.

        Returns:
            dict with keys:
                compliance_score          — float 0–100
                category_scores           — {numbering, spacing, marks, header, indent}
                issues_found              — list of issue strings
                edge_cases                — list of edge-case descriptions
                math_expressions_preserved — bool
                summary                   — one-sentence summary
                fallback_mode             — always False on success

        On failure:
            {compliance_score: None, fallback_mode: True, error: str}
        """
        try:
            return await self._call_llm(original, formatted, template_rules)
        except Exception as exc:
            logger.warning("LLMValidator: falling back to rule-based only — %s", exc)
            return {
                "compliance_score": None,
                "category_scores": {},
                "issues_found": [],
                "edge_cases": [],
                "math_expressions_preserved": True,
                "summary": "LLM validation unavailable.",
                "fallback_mode": True,
                "error": str(exc),
            }

    async def _call_llm(
        self,
        original: Document,
        formatted: Document,
        template_rules: dict,
    ) -> dict:
        """Internal: build prompt and call GPT-4o-mini."""
        from azure.ai.projects import AIProjectClient
        from azure.identity import DefaultAzureCredential

        original_text = "\n".join(p.text for p in original.paragraphs)
        formatted_text = "\n".join(p.text for p in formatted.paragraphs)

        user_prompt = (
            f"TEMPLATE RULES:\n{json.dumps(template_rules, indent=2)}\n\n"
            f"ORIGINAL DOCUMENT:\n{original_text[:4000]}\n\n"
            f"FORMATTED DOCUMENT:\n{formatted_text[:4000]}"
        )

        client = AIProjectClient(
            endpoint=self._endpoint,
            credential=DefaultAzureCredential(),
        )

        response = client.inference.get_chat_completions(
            model=os.getenv("AZURE_FOUNDRY_DEPLOYMENT", "gpt-4o-mini"),
            messages=[
                {"role": "system", "content": self.SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
            timeout=30,
        )

        raw = response.choices[0].message.content.strip()
        result = json.loads(raw)
        result["fallback_mode"] = False
        return result


# ── Thin Wrapper — FormattingEngineAgent ────────────────────────────────────


class FormattingEngineAgent:
    """
    Thin orchestration wrapper that runs Layer 1 then Layer 2.

    Usage::

        agent = FormattingEngineAgent()
        formatted_doc, validation = await agent.process_and_validate(
            original_doc, template_rules
        )
    """

    def __init__(self) -> None:
        self.rule_based = RuleBasedFormatter()
        self.llm_validator = LLMValidator()

    async def process_and_validate(
        self, original: Document, template_rules: dict
    ) -> tuple:
        """
        Run Layer 1 (rule-based) then Layer 2 (LLM validation).

        Args:
            original:       Original python-docx Document.
            template_rules: Template rules dict from Azure AI Search.

        Returns:
            Tuple of (formatted_doc: Document, validation_result: dict).
            If LLM validation fails, validation_result contains
            ``fallback_mode=True`` and ``compliance_score=None``.
        """
        # Layer 1 — deterministic transforms
        formatted_doc = self.rule_based.process(original, template_rules)

        # Layer 2 — LLM validation (never raises; returns fallback on error)
        validation_result = await self.llm_validator.validate(
            original, formatted_doc, template_rules
        )

        return formatted_doc, validation_result
