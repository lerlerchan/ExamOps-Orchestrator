"""
scripts/upload_template.py
==========================
Parse sample/ guideline .docx files, extract structured template rules, generate
an ada-002 embedding, and upsert into the Azure AI Search ``exam-templates`` index.

Usage
-----
    python scripts/upload_template.py                  # upload to Azure AI Search
    python scripts/upload_template.py --dry-run        # print rules, skip upload
    python scripts/upload_template.py --index-name my-index

Environment variables required for upload (not needed with --dry-run):
    AZURE_SEARCH_ENDPOINT
    AZURE_SEARCH_KEY
    AZURE_OPENAI_ENDPOINT
    AZURE_OPENAI_KEY
    AZURE_EMBEDDING_MODEL   (optional, default: text-embedding-ada-002)

The index must already exist with the schema described in AZURE_SETUP.md.
Run this script once per template change to keep the vector store up to date.
"""

import argparse
import json
import logging
import os
import re
import sys
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# ── Paths ────────────────────────────────────────────────────────────────────

REPO_ROOT = Path(__file__).resolve().parent.parent
SAMPLE_DIR = REPO_ROOT / "sample"

GUIDELINE_FILES = {
    "exam_paper_format": "Exam Paper Format Guideline 1.docx",
    "marking_scheme_format": "Marking Scheme Format Guideline 1.docx",
}


# ── Rule extraction ───────────────────────────────────────────────────────────


def _extract_text(doc: Document) -> list[str]:
    """Return non-empty paragraph texts from a Document."""
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _detect_numbering_scheme(lines: list[str]) -> list[str]:
    """
    Infer the numbering hierarchy from the document's paragraph texts.

    Returns a list of format strings, e.g. ["Q{n}.", "(a)", "(i)"].
    Falls back to the institutional default if detection fails.
    """
    scheme = []
    for line in lines:
        if re.match(r"^Q\d+\.", line) and "Q{n}." not in scheme:
            scheme.append("Q{n}.")
        elif re.match(r"^\([a-z]\)", line) and "(a)" not in scheme:
            scheme.append("(a)")
        elif re.match(r"^\([ivxlcdm]+\)", line, re.IGNORECASE) and "(i)" not in scheme:
            scheme.append("(i)")
    return scheme or ["Q{n}.", "(a)", "(i)"]


def _detect_marks_pattern(lines: list[str]) -> str:
    """
    Return a regex string for marks notation found in the document.

    Prefers ``(N marks)``; falls back to ``[N marks]`` if not found.
    """
    for line in lines:
        if re.search(r"\(\d+\s*marks?\)", line, re.IGNORECASE):
            return r"(\d+ marks?)"
        if re.search(r"\[\d+\s*marks?\]", line, re.IGNORECASE):
            return r"[\d+ marks?]"
    return r"(\d+ marks?)"


def _detect_header(doc: Document) -> str:
    """Return the first-section header text, or the institutional default."""
    try:
        header_paras = doc.sections[0].header.paragraphs
        text = " ".join(p.text for p in header_paras).strip()
        if text:
            return text
    except Exception:
        pass
    # Fall back to first non-empty paragraph that looks like an institution name
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and len(t) < 80 and t == t.upper():
            return t
    return "SOUTHERN UNIVERSITY COLLEGE"


def extract_template_rules(doc: Document, title: str) -> dict:
    """
    Extract structured template rules from a guideline Document.

    Args:
        doc:   python-docx Document.
        title: Human-readable title used as the ``title`` field in the index.

    Returns:
        dict with keys:
            id                  — slug derived from title
            title               — human-readable name
            header_text         — first-section header string
            footer_text         — first-section footer string (may be empty)
            numbering_scheme    — list of format strings
            marks_pattern       — regex string for marks notation
            colon_spacing       — True (always enforced)
            margin_cm           — {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.5}
            indentation_cm      — {"l1": 0.0, "l2": 1.5, "l3": 3.0}
            content             — JSON string of the rules (used for embedding)
    """
    lines = _extract_text(doc)

    # Header / footer
    header_text = _detect_header(doc)
    try:
        footer_paras = doc.sections[0].footer.paragraphs
        footer_text = " ".join(p.text for p in footer_paras).strip()
    except Exception:
        footer_text = ""

    # Detect section margins from the first section
    margin_cm = {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.5}
    try:
        sec = doc.sections[0]
        # python-docx stores margins in EMU; 914400 EMU = 1 inch = 2.54 cm
        emu_per_cm = 360000
        margin_cm = {
            "top": round(sec.top_margin / emu_per_cm, 2),
            "bottom": round(sec.bottom_margin / emu_per_cm, 2),
            "left": round(sec.left_margin / emu_per_cm, 2),
            "right": round(sec.right_margin / emu_per_cm, 2),
        }
    except Exception:
        pass

    rules = {
        "header_text": header_text,
        "footer_text": footer_text,
        "numbering_scheme": _detect_numbering_scheme(lines),
        "marks_pattern": _detect_marks_pattern(lines),
        "colon_spacing": True,
        "margin_cm": margin_cm,
        "indentation_cm": {"l1": 0.0, "l2": 1.5, "l3": 3.0},
    }

    doc_id = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")
    content_json = json.dumps(rules, ensure_ascii=False)

    return {
        "id": doc_id,
        "title": title,
        "content": content_json,
        # template_rules is the field CoordinatorAgent reads back from the index
        "template_rules": rules,
        **rules,
    }


# ── Embedding + upload ────────────────────────────────────────────────────────


def _generate_embedding(text: str) -> list[float]:
    """
    Generate an ada-002 embedding for ``text`` via Azure OpenAI.

    TODO: swap to managed identity when AZURE_OPENAI_KEY is removed from env.
    """
    from openai import AzureOpenAI

    client = AzureOpenAI(
        azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
        api_key=os.environ["AZURE_OPENAI_KEY"],
        api_version="2024-02-01",
    )
    model = os.getenv("AZURE_EMBEDDING_MODEL", "text-embedding-ada-002")
    response = client.embeddings.create(input=text, model=model)
    return response.data[0].embedding


def _upsert_document(doc_dict: dict, index_name: str) -> None:
    """
    Upsert a single document (with its embedding) into Azure AI Search.

    The document must contain ``id``, ``title``, ``content``,
    ``content_vector``, and ``template_rules`` fields matching the
    index schema defined in AZURE_SETUP.md.
    """
    from azure.search.documents import SearchClient
    from azure.core.credentials import AzureKeyCredential

    client = SearchClient(
        endpoint=os.environ["AZURE_SEARCH_ENDPOINT"],
        index_name=index_name,
        credential=AzureKeyCredential(os.environ["AZURE_SEARCH_KEY"]),
    )
    result = client.upload_documents(documents=[doc_dict])
    for r in result:
        if r.succeeded:
            logger.info("Upserted document id=%s into index '%s'", r.key, index_name)
        else:
            logger.error(
                "Failed to upsert id=%s: %s", r.key, r.errors
            )


# ── CLI entry point ───────────────────────────────────────────────────────────


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Upload template rules from sample/ guidelines into Azure AI Search."
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Extract and print template rules without uploading.",
    )
    parser.add_argument(
        "--index-name",
        default=os.getenv("SEARCH_INDEX_NAME", "exam-templates"),
        help="Azure AI Search index name (default: exam-templates).",
    )
    args = parser.parse_args()

    if not SAMPLE_DIR.exists():
        logger.error("sample/ directory not found at %s", SAMPLE_DIR)
        sys.exit(1)

    processed = 0
    for template_key, filename in GUIDELINE_FILES.items():
        filepath = SAMPLE_DIR / filename
        if not filepath.exists():
            logger.warning("Skipping %s — file not found", filepath)
            continue

        logger.info("Processing %s ...", filename)
        doc = Document(filepath)
        title = filepath.stem  # e.g. "Exam Paper Format Guideline 1"
        rules = extract_template_rules(doc, title)

        if args.dry_run:
            print(f"\n{'=' * 60}")
            print(f"File   : {filename}")
            print(f"Title  : {rules['title']}")
            print(f"Header : {rules['header_text']}")
            print(f"Margins: {rules['margin_cm']}")
            print(f"Numbering: {rules['numbering_scheme']}")
            print(f"Marks pattern: {rules['marks_pattern']}")
            print(f"Content JSON:\n{rules['content']}")
            processed += 1
            continue

        # Generate embedding from the content JSON
        logger.info("Generating embedding for '%s' ...", title)
        embedding = _generate_embedding(rules["content"])

        upload_doc = {
            "id": rules["id"],
            "title": rules["title"],
            "content": rules["content"],
            "content_vector": embedding,
            # Store full rules dict as serialized JSON so search results
            # can be deserialized by FileHandlerAgent.get_template_from_vectordb()
            "template_rules": json.dumps(rules["template_rules"], ensure_ascii=False),
        }

        _upsert_document(upload_doc, args.index_name)
        processed += 1

    if processed == 0:
        logger.error("No guideline files were found in %s", SAMPLE_DIR)
        sys.exit(1)

    verb = "Printed" if args.dry_run else "Uploaded"
    logger.info("%s %d template rule document(s).", verb, processed)


if __name__ == "__main__":
    main()
